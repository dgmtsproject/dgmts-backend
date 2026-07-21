"""Generate DGMTS_DEVELOPER_GUIDE.docx — single consolidated developer documentation."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OUTPUT = Path(__file__).resolve().parent.parent.parent / "DGMTS_DEVELOPER_GUIDE.docx"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p


def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def add_numbered(doc, text):
    doc.add_paragraph(text, style="List Number")


def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)


def build():
    doc = Document()
    doc.core_properties.title = "DGMTS Developer Guide"
    doc.core_properties.subject = "dgmts + dgmts-backend internal documentation"

    # Title page
    t = doc.add_heading("DGMTS Developer Guide", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Instrumentation monitoring platform — frontend, backend, Supabase, VPS deployment")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Repositories: dgmts (React) + dgmts-backend (Flask)")
    doc.add_paragraph("Confidential — internal use only")
    doc.add_page_break()

    # 1. Overview
    add_heading(doc, "1. System overview", 1)
    add_para(doc, "DGMTS is a geotechnical instrumentation monitoring platform. Clients view live and historical sensor data (seismographs, tiltmeters, Instantel/Micromate, AMTS), manage projects and instruments, and receive email alerts when thresholds are exceeded.")
    add_heading(doc, "1.1 Two repositories", 2)
    add_bullet(doc, "dgmts — React + TypeScript SPA (Vite). UI, graphs, Supabase CRUD for projects/instruments/users.")
    add_bullet(doc, "dgmts-backend — Flask API on port 5000, background alert scheduler, email sending, external API integration.")
    add_heading(doc, "1.2 Production URLs", 2)
    add_bullet(doc, "Backend API: https://imsite.dullesgeotechnical.com")
    add_bullet(doc, "Frontend: deployed via Vercel (see vercel.json in dgmts)")
    add_bullet(doc, "Database: Supabase (instrumentation tables)")
    add_heading(doc, "1.3 Architecture", 2)
    add_para(doc, "The React frontend talks to Supabase directly for most CRUD (projects, instruments, alarms list). The Flask backend owns scheduled alert checks, email delivery, authentication, payments, and proxy endpoints for Syscom/Micromate/sensor APIs.")
    add_code_block(doc, """Flow:
  Browser (dgmts React) ──► Supabase (Projects, instruments, sent_alerts, …)
  Browser ──► imsite Flask API (auth, payments, some proxies)
  Flask scheduler (every minute) ──► Supabase + Syscom/Micromate APIs ──► SMTP emails""")

    # 2. VPS
    add_heading(doc, "2. VPS deployment (backend)", 1)
    add_heading(doc, "2.1 Server access", 2)
    add_bullet(doc, "Host: 192.168.1.219")
    add_bullet(doc, "User: root")
    add_bullet(doc, "Password: ZFatima@2025")
    add_code_block(doc, "ssh root@192.168.1.219")
    add_heading(doc, "2.2 tmux session", 2)
    add_bullet(doc, "Session name: flaskserver")
    add_bullet(doc, "Attach: tmux attach -t flaskserver")
    add_bullet(doc, "Detach (keep server running): press Ctrl+b, then d")
    add_bullet(doc, "Exit SSH: type exit and press Enter")
    add_heading(doc, "2.3 Deploy / restart procedure", 2)
    add_numbered(doc, "SSH into the VPS: ssh root@192.168.1.219")
    add_numbered(doc, "Attach to tmux: tmux attach -t flaskserver")
    add_numbered(doc, "Go to backend directory (confirm with pwd on server): cd /path/to/dgmts-backend")
    add_numbered(doc, "Pull changes: git pull origin main")
    add_numbered(doc, "Activate venv: source venv/bin/activate")
    add_numbered(doc, "If requirements.txt changed: pip install -r requirements.txt")
    add_numbered(doc, "Stop old gunicorn with Ctrl+C in the tmux pane")
    add_numbered(doc, "Start server:")
    add_code_block(doc, "gunicorn app:app --certfile=cert.pem --keyfile=key.pem -b 0.0.0.0:5000")
    add_numbered(doc, "Detach from tmux: Ctrl+b, then d")
    add_numbered(doc, "Exit SSH: exit")
    add_heading(doc, "2.4 tmux quick reference", 2)
    add_bullet(doc, "List sessions: tmux ls")
    add_bullet(doc, "Scroll log output: Ctrl+b, then [ — press q to exit scroll mode")
    add_bullet(doc, "New session (if needed): tmux new -s flaskserver")
    add_heading(doc, "2.5 Verify deployment", 2)
    add_code_block(doc, "curl -k https://localhost:5000/\ncurl https://imsite.dullesgeotechnical.com/")
    add_para(doc, "Expected JSON: message DGMTS Backend API, status running.")

    # 3. Frontend
    add_heading(doc, "3. dgmts (frontend)", 1)
    add_heading(doc, "3.1 Stack", 2)
    add_bullet(doc, "React 19, TypeScript, Vite 6, MUI, Plotly, Leaflet, Supabase JS client")
    add_heading(doc, "3.2 Key files", 2)
    add_bullet(doc, "src/supabase.ts — Supabase client")
    add_bullet(doc, "src/config.ts — API_BASE_URL (https://imsite.dullesgeotechnical.com)")
    add_bullet(doc, "src/Routes/Routes.tsx — all routes")
    add_bullet(doc, "src/context/AdminContext.tsx — auth via /api/check-auth")
    add_bullet(doc, "vercel.json — production API rewrites to Syscom and imsite")
    add_heading(doc, "3.3 Environment variables (.env — do not commit)", 2)
    add_code_block(doc, """VITE_SUPABASE_PROJECT_URL=https://xxxxx.supabase.co
VITE_SUPABASE_ANON_KEY=eyJ...
VITE_SYSCOM_API_KEY=your-syscom-key
VITE_MAPBOX_ACCESS_TOKEN=pk....""")
    add_heading(doc, "3.4 Local development", 2)
    add_code_block(doc, "cd dgmts\nnpm install\nnpm run dev")
    add_para(doc, "For local backend: set API_BASE_URL to http://localhost:5000 in src/config.ts")
    add_heading(doc, "3.5 Production build", 2)
    add_code_block(doc, "npm run build   # output: dist/\n# Deploy via Vercel; set VITE_* env vars in hosting dashboard")

    # 4. Backend
    add_heading(doc, "4. dgmts-backend", 1)
    add_heading(doc, "4.1 Stack", 2)
    add_bullet(doc, "Flask + flask-cors, Supabase Python client, schedule (background jobs), gunicorn (production)")
    add_heading(doc, "4.2 Project layout", 2)
    add_bullet(doc, "app.py — Flask entry, blueprints, starts scheduler on import")
    add_bullet(doc, "config.py — environment configuration")
    add_bullet(doc, "routes/ — auth, email, sensor, micromate, payment, static site APIs")
    add_bullet(doc, "services/ — alert_service, duration_alert_service, micromate, rock seismograph, email")
    add_bullet(doc, "utils/scheduler.py — minute-by-minute job registration")
    add_bullet(doc, "sql/ — Supabase migration scripts")
    add_heading(doc, "4.3 Environment variables (.env on server)", 2)
    add_code_block(doc, """FLASK_SECRET_KEY=...
EMAIL_USERNAME=...
EMAIL_PASSWORD=...
SUPABASE_URL=...
SUPABASE_KEY=...
SYSCOM_API_KEY=...
# See env.example for full list""")
    add_heading(doc, "4.4 Local development", 2)
    add_code_block(doc, """cd dgmts-backend
python -m venv venv
venv\\Scripts\\activate          # Windows
pip install -r requirements.txt
python app.py                     # :5000, scheduler auto-starts""")

    # 5. Scheduler
    add_heading(doc, "5. Background scheduler", 1)
    add_para(doc, "Runs in a daemon thread when app.py loads (including under gunicorn). Restart gunicorn after any scheduler.py change.")
    add_bullet(doc, "Every minute: SMG-1 seismograph alerts")
    add_bullet(doc, "Every minute: SMG-3 seismograph alerts")
    add_bullet(doc, "Every minute: Instrument 13453 alerts")
    add_bullet(doc, "Every minute: Instantel 1 (Micromate) alerts")
    add_bullet(doc, "Every minute: Instantel 2 (UM16368) alerts")
    add_bullet(doc, "Every minute: Duration-based alerts (all configured instruments)")
    add_para(doc, "Disabled by default: tiltmeter data fetch, ROCKSMG alerts (commented in scheduler.py).")

    # 6. Database
    add_heading(doc, "6. Supabase database", 1)
    add_heading(doc, "6.1 Main tables", 2)
    add_bullet(doc, "Projects — project_id_second for editable display ID; id is immutable PK")
    add_bullet(doc, "instruments — thresholds, emails, syscom_device_id, duration config")
    add_bullet(doc, "sent_alerts — alert deduplication + Alarms UI")
    add_bullet(doc, "sent_alert_logs — audit trail (backend logging; FK to instrument_id)")
    add_bullet(doc, "sensor_readings — tiltmeter history")
    add_heading(doc, "6.2 Migration scripts (run in Supabase SQL Editor)", 2)
    add_bullet(doc, "sql/add_instrument_id_second.sql — editable instrument display ID")
    add_bullet(doc, "sql/add_duration_alert_columns.sql — duration alert fields")
    add_bullet(doc, "sql/create_about_employees.sql — static site employees (if used)")
    add_heading(doc, "6.3 Safe migration order", 2)
    add_numbered(doc, "Write idempotent SQL (ADD COLUMN IF NOT EXISTS, backfill, then constraints)")
    add_numbered(doc, "Run in Supabase SQL Editor")
    add_numbered(doc, "Deploy backend (if Python reads new columns)")
    add_numbered(doc, "Deploy frontend (if UI edits new fields)")

    # 7. Domain patterns
    add_heading(doc, "7. Important domain patterns", 1)
    add_heading(doc, "7.1 Editable IDs without breaking FKs", 2)
    add_para(doc, "Projects: keep Projects.id immutable; users edit project_id_second (shown in UI).")
    add_para(doc, "Instruments: keep instrument_id immutable (referenced by sent_alert_logs, sent_alerts, backend jobs); users edit instrument_id_second.")
    add_para(doc, "Never UPDATE instrument_id or Projects.id when child rows exist — you will get foreign key errors.")
    add_heading(doc, "7.2 Alert types", 2)
    add_para(doc, "Instant alerts: fire when any single reading crosses alert_value / warning_value / shutdown_value.")
    add_para(doc, "Duration alerts: fire when threshold is continuously exceeded for duration_seconds. Separate thresholds and email lists. alert_type in sent_alerts: duration_alert, duration_warning, duration_shutdown.")
    add_heading(doc, "7.3 Instrument configuration UI", 2)
    add_bullet(doc, "Edit Instrument — tab Instant Threshold Alerts + tab Duration Achieved Alerts")
    add_bullet(doc, "Edit Tiltmeter Instrument — same tabs with X/Y/Z JSON thresholds")
    add_bullet(doc, "Instruments List — shows instant and duration columns")

    # 8. API endpoints
    add_heading(doc, "8. Key API endpoints", 1)
    add_bullet(doc, "GET / — health check")
    add_bullet(doc, "POST /api/login, GET /api/check-auth, POST /api/logout")
    add_bullet(doc, "POST /api/forgot-password, POST /api/reset-password")
    add_bullet(doc, "GET /api/sensor-data/<node_id>")
    add_bullet(doc, "GET /api/micromate/readings — Instantel 1")
    add_bullet(doc, "GET /api/micromate/UM16368/readings — Instantel 2")
    add_bullet(doc, "Test routes in routes/email_routes.py — manual alert triggers")

    # 9. Upgrade-safe guidelines
    add_heading(doc, "9. Upgrade-safe development guidelines", 1)
    add_heading(doc, "9.1 Do", 2)
    add_bullet(doc, "Use project_id_second and instrument_id_second for user-visible ID changes")
    add_bullet(doc, "Add new config as new columns with defaults")
    add_bullet(doc, "Use distinct alert_type values for new alert kinds")
    add_bullet(doc, "Test with /api/test-* routes before relying on scheduler")
    add_bullet(doc, "Deploy order: SQL → backend restart → frontend")
    add_heading(doc, "9.2 Avoid", 2)
    add_bullet(doc, "Changing PK instrument_id or Projects.id when FK rows exist")
    add_bullet(doc, "Deleting sent_alert_logs to fix update errors")
    add_bullet(doc, "Running multiple gunicorn instances on port 5000")
    add_bullet(doc, "Committing .env or passwords to git")

    # 10. Troubleshooting
    add_heading(doc, "10. Troubleshooting", 1)
    add_bullet(doc, "Alerts not sending — tmux running? gunicorn up? EMAIL_* in .env? emails on instrument row?")
    add_bullet(doc, "FK error updating instrument — use instrument_id_second only")
    add_bullet(doc, "Frontend auth loops — check API_BASE_URL, CORS, FLASK_SECRET_KEY")
    add_bullet(doc, "Graphs empty — VITE_SYSCOM_API_KEY, syscom_device_id on instrument")
    add_bullet(doc, "Duration alerts never fire — duration_seconds + thresholds + emails set? enough readings in window?")
    add_bullet(doc, "Scheduler changes not applied — must restart gunicorn")

    # 11. Maintenance scripts
    add_heading(doc, "11. Maintenance scripts (on VPS with venv active)", 1)
    add_code_block(doc, """python delete_instantel_info_logs.py
python send_missed_smg1_alerts.py
python send_missed_rock_seismograph_alerts.py""")

    doc.add_page_break()
    add_heading(doc, "Document revision", 1)
    add_para(doc, "Update this document when adding scheduler jobs, Supabase columns, VPS layout, or production URLs.")
    add_para(doc, "Generated for dgmts + dgmts-backend teams.")

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    build()
